"""
Export Service
==============

Convert markdown content to .docx or .pdf.

Dependencies:
  pip install python-docx markdown weasyprint
"""

from __future__ import annotations

import re
from io import BytesIO


# ── DOCX ──────────────────────────────────────────────────────────────────────

def markdown_to_docx(content: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Remove default empty paragraph that python-docx adds
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ", "* ", "+ ")):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline(para, line[2:])
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            para = doc.add_paragraph(style="List Number")
            _add_inline(para, text)
        elif re.match(r"^(-{3,}|_{3,}|\*{3,})$", line):
            doc.add_paragraph("─" * 40)
        elif line == "":
            # Blank line between paragraphs — only add space if prev wasn't a heading
            pass
        else:
            para = doc.add_paragraph()
            _add_inline(para, line)

        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_inline(para, text: str) -> None:
    """Parse **bold**, *italic*, and plain text into runs."""
    # Pattern matches **bold**, *italic*, or plain text segments
    pattern = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)")
    for match in pattern.finditer(text):
        bold_text, italic_text, plain_text = match.groups()
        if bold_text:
            run = para.add_run(bold_text)
            run.bold = True
        elif italic_text:
            run = para.add_run(italic_text)
            run.italic = True
        elif plain_text:
            para.add_run(plain_text)


# ── PDF ───────────────────────────────────────────────────────────────────────

_PDF_CSS = """
body {
    font-family: Georgia, 'Times New Roman', serif;
    font-size: 11pt;
    line-height: 1.65;
    color: #1a1a1a;
    max-width: 740px;
    margin: 48px auto;
    padding: 0 24px;
}
h1 { font-size: 20pt; margin-bottom: 4px; font-weight: 700; }
h2 { font-size: 14pt; margin-top: 22px; margin-bottom: 4px; border-bottom: 1px solid #ccc; padding-bottom: 2px; }
h3 { font-size: 12pt; margin-top: 16px; margin-bottom: 2px; }
p  { margin: 6px 0 10px; }
ul, ol { margin: 4px 0 10px 20px; }
li { margin-bottom: 3px; }
hr { border: none; border-top: 1px solid #ccc; margin: 18px 0; }
strong { font-weight: 700; }
em { font-style: italic; }
"""


def markdown_to_pdf(content: str) -> bytes:
    import markdown as md
    try:
        from weasyprint import HTML, CSS
    except Exception as e:
        if "libgobject-2.0-0" in str(e) or "cannot load library" in str(e).lower():
            raise ImportError(
                "WeasyPrint dependency missing: libgobject-2.0-0. "
                "On Windows, you must install the GTK for Windows Runtime: "
                "https://github.com/tschoonj/GTK-for-Windows-Runtime-Installer/releases"
            ) from e
        raise

    html_body = md.markdown(
        content,
        extensions=["tables", "fenced_code", "nl2br"],
    )
    full_html = (
        "<!DOCTYPE html><html><head><meta charset='utf-8'></head>"
        f"<body>{html_body}</body></html>"
    )
    return HTML(string=full_html).write_pdf(stylesheets=[CSS(string=_PDF_CSS)])
